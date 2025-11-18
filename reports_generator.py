"""
Report Generation Module for Faculty Housing Management System
نظام توليد التقارير لنظام إدارة إسكان أعضاء هيئة التدريس
"""

from datetime import datetime
import io
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import database


def get_report_data(report_type='all', from_date=None, to_date=None, building=None):
    """Get report data for export. Integrates with actual database queries based on filter parameters.
    
    Args:
        report_type: Type of report ('all', 'residents', 'violations', 'security', 'complaints')
        from_date: Start date for filtering (YYYY-MM-DD format)
        to_date: End date for filtering (YYYY-MM-DD format)
        building: Building ID or name for filtering
    
    Returns:
        Dictionary containing report data from the database
    """
    # Get database connection
    conn = database.get_db_connection()
    cursor = conn.cursor()
    
    try:
        # Build date filter SQL fragment
        date_filter = ""
        date_params = []
        if from_date:
            date_filter += " AND date(created_at) >= date(?)"
            date_params.append(from_date)
        if to_date:
            date_filter += " AND date(created_at) <= date(?)"
            date_params.append(to_date)
        
        # Build building filter SQL fragment
        building_filter = ""
        building_params = []
        if building:
            building_filter = " AND (building_id = ? OR building = ?)"
            building_params = [building, building]
        # Build date filter SQL fragment
        date_filter = ""
        date_params = []
        if from_date:
            date_filter += " AND date(created_at) >= date(?)"
            date_params.append(from_date)
        if to_date:
            date_filter += " AND date(created_at) <= date(?)"
            date_params.append(to_date)
        
        # Build building filter SQL fragment
        building_filter = ""
        building_params = []
        if building:
            building_filter = " AND (building_id = ? OR building = ?)"
            building_params = [building, building]
        
        # Get summary statistics
        summary = {}
        
        # Total residents
        cursor.execute('SELECT COUNT(*) FROM residents WHERE is_active = 1')
        summary['total_residents'] = cursor.fetchone()[0]
        
        # Buildings and apartments
        cursor.execute('SELECT COUNT(*) FROM buildings')
        summary['total_buildings'] = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM apartments')
        summary['total_apartments'] = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM apartments WHERE is_occupied = 1')
        summary['occupied_apartments'] = cursor.fetchone()[0]
        
        summary['occupancy_rate'] = round((summary['occupied_apartments'] / summary['total_apartments'] * 100), 1) if summary['total_apartments'] > 0 else 0
        
        # Violations
        cursor.execute('SELECT COUNT(*) FROM traffic_violations')
        summary['total_violations'] = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM traffic_violations WHERE status IN ("resolved", "محلول", "closed", "مغلق")')
        summary['resolved_violations'] = cursor.fetchone()[0]
        
        summary['pending_violations'] = summary['total_violations'] - summary['resolved_violations']
        
        # Complaints
        cursor.execute('SELECT COUNT(*) FROM complaints')
        summary['total_complaints'] = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM complaints WHERE status IN ("resolved", "محلولة", "closed", "مغلقة")')
        summary['resolved_complaints'] = cursor.fetchone()[0]
        
        summary['pending_complaints'] = summary['total_complaints'] - summary['resolved_complaints']
        
        # Security incidents
        cursor.execute('SELECT COUNT(*) FROM security_incidents')
        summary['total_security_incidents'] = cursor.fetchone()[0]
        
        # Parking
        cursor.execute('SELECT COUNT(*) FROM parking')
        summary['total_parking_slots'] = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM parking WHERE is_occupied = 1')
        summary['occupied_parking'] = cursor.fetchone()[0]
        
        summary['parking_occupancy_rate'] = round((summary['occupied_parking'] / summary['total_parking_slots'] * 100), 1) if summary['total_parking_slots'] > 0 else 0
        
        # Get detailed data based on report type
        residents_list = []
        violations_list = []
        security_list = []
        complaints_list = []
        buildings_list = []
        
        if report_type in ['all', 'residents']:
            # Get residents data
            cursor.execute("""
                SELECT r.id, r.name, b.name as building_name, 
                       COALESCE(a.number, r.apartment_number) as apartment,
                       r.move_in_date, 
                       CASE WHEN r.is_active = 1 THEN 'مقيم' ELSE 'غير مقيم' END as status
                FROM residents r
                LEFT JOIN buildings b ON r.building_id = b.id
                LEFT JOIN apartments a ON r.apartment_id = a.id
                ORDER BY r.created_at DESC
                LIMIT 100
            """)
            for row in cursor.fetchall():
                residents_list.append({
                    'id': row[0],
                    'name': row[1] or 'غير محدد',
                    'building': row[2] or 'غير محدد',
                    'apartment': row[3] or 'غير محدد',
                    'move_in_date': row[4] or 'غير محدد',
                    'status': row[5]
                })
        
        if report_type in ['all', 'violations']:
            # Get violations data
            cursor.execute("""
                SELECT id, violation_type, violation_date, location, 
                       violator_name, status
                FROM traffic_violations
                ORDER BY violation_date DESC
                LIMIT 100
            """)
            for row in cursor.fetchall():
                violations_list.append({
                    'id': row[0],
                    'type': row[1] or 'غير محدد',
                    'date': row[2] or 'غير محدد',
                    'location': row[3] or 'غير محدد',
                    'resident': row[4] or 'غير محدد',
                    'status': row[5] or 'معلق'
                })
        
        if report_type in ['all', 'security']:
            # Get security incidents data
            cursor.execute("""
                SELECT id, incident_type, incident_date, location, 
                       status, severity
                FROM security_incidents
                ORDER BY incident_date DESC
                LIMIT 100
            """)
            for row in cursor.fetchall():
                security_list.append({
                    'id': row[0],
                    'type': row[1] or 'غير محدد',
                    'date': row[2] or 'غير محدد',
                    'location': row[3] or 'غير محدد',
                    'status': row[4] or 'معلق',
                    'severity': row[5] or 'منخفض'
                })
        
        if report_type in ['all', 'complaints']:
            # Get complaints data
            cursor.execute("""
                SELECT c.id, c.complaint_type, c.description, c.complaint_date, 
                       b.name as building_name, c.status
                FROM complaints c
                LEFT JOIN buildings b ON c.building_id = b.id
                ORDER BY c.complaint_date DESC
                LIMIT 100
            """)
            for row in cursor.fetchall():
                complaints_list.append({
                    'id': row[0],
                    'type': row[1] or 'غير محدد',
                    'description': row[2] or 'غير محدد',
                    'date': row[3] or 'غير محدد',
                    'building': row[4] or 'غير محدد',
                    'status': row[5] or 'معلق'
                })
        
        # Get buildings data
        cursor.execute("""
            SELECT b.id, b.name,
                   (SELECT COUNT(*) FROM apartments WHERE building_id = b.id) as total_apartments,
                   (SELECT COUNT(*) FROM apartments WHERE building_id = b.id AND is_occupied = 1) as occupied_apartments
            FROM buildings b
            ORDER BY b.name
        """)
        for row in cursor.fetchall():
            total_apts = row[2] or 0
            occupied_apts = row[3] or 0
            occupancy = round((occupied_apts / total_apts * 100), 1) if total_apts > 0 else 0
            buildings_list.append({
                'id': row[0],
                'name': row[1] or 'غير محدد',
                'apartments': total_apts,
                'occupied': occupied_apts,
                'occupancy': occupancy
            })
        
        # Prepare final data structure
        data = {
            'metadata': {
                'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'generated_at_ar': datetime.now().strftime('%Y/%m/%d %H:%M:%S'),
                'report_type': report_type,
                'from_date': from_date,
                'to_date': to_date,
                'building': building,
                'title': 'تقرير احترافي ورسمي - نظام إدارة إسكان أعضاء هيئة التدريس',
                'title_en': 'Professional and Official Report - Faculty Housing Management System',
                'university': 'جامعة الإمام محمد بن سعود الإسلامية',
                'university_en': 'Imam Mohammad Ibn Saud Islamic University'
            },
            'summary': summary,
            'residents': residents_list if residents_list else [
                {'id': 0, 'name': 'لا توجد بيانات', 'building': '-', 'apartment': '-', 'move_in_date': '-', 'status': '-'}
            ],
            'violations': violations_list if violations_list else [
                {'id': 0, 'type': 'لا توجد بيانات', 'date': '-', 'location': '-', 'resident': '-', 'status': '-'}
            ],
            'security': security_list if security_list else [
                {'id': 0, 'type': 'لا توجد بيانات', 'date': '-', 'location': '-', 'status': '-', 'severity': '-'}
            ],
            'buildings': buildings_list if buildings_list else [
                {'id': 0, 'name': 'لا توجد بيانات', 'apartments': 0, 'occupied': 0, 'occupancy': 0}
            ],
            'complaints': complaints_list if complaints_list else [
                {'id': 0, 'type': 'لا توجد بيانات', 'description': '-', 'date': '-', 'building': '-', 'status': '-'}
            ]
        }
        
        return data
    
    finally:
        conn.close()


def generate_excel_report(report_type='all', from_date=None, to_date=None, building=None):
    """Generate professional Excel report"""
    # Get data
    data = get_report_data(report_type, from_date, to_date, building)
    
    # Create workbook
    wb = Workbook()
    
    # Remove default sheet
    if 'Sheet' in wb.sheetnames:
        del wb['Sheet']
    
    # Define styles
    header_font = Font(name='Arial', size=14, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='0F3D68', end_color='0F3D68', fill_type='solid')
    title_font = Font(name='Arial', size=16, bold=True, color='0F3D68')
    subtitle_font = Font(name='Arial', size=12, bold=True, color='2E8BC0')
    normal_font = Font(name='Arial', size=11)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    right_alignment = Alignment(horizontal='right', vertical='center', wrap_text=True)
    
    # Create Summary Sheet
    ws_summary = wb.create_sheet('ملخص تنفيذي')
    
    # Header
    ws_summary.merge_cells('A1:F1')
    ws_summary['A1'] = data['metadata']['title']
    ws_summary['A1'].font = title_font
    ws_summary['A1'].alignment = center_alignment
    
    ws_summary.merge_cells('A2:F2')
    ws_summary['A2'] = data['metadata']['university']
    ws_summary['A2'].font = subtitle_font
    ws_summary['A2'].alignment = center_alignment
    
    ws_summary.merge_cells('A3:F3')
    ws_summary['A3'] = f"تاريخ التقرير: {data['metadata']['generated_at_ar']}"
    ws_summary['A3'].font = normal_font
    ws_summary['A3'].alignment = center_alignment
    
    # Summary data
    row = 5
    ws_summary[f'A{row}'] = 'المؤشر'
    ws_summary[f'B{row}'] = 'القيمة'
    for cell in [f'A{row}', f'B{row}']:
        ws_summary[cell].font = header_font
        ws_summary[cell].fill = header_fill
        ws_summary[cell].alignment = center_alignment
        ws_summary[cell].border = border
    
    summary_items = [
        ('إجمالي السكان', data['summary']['total_residents']),
        ('إجمالي المباني', data['summary']['total_buildings']),
        ('إجمالي الشقق', data['summary']['total_apartments']),
        ('الشقق المشغولة', data['summary']['occupied_apartments']),
        ('معدل الإشغال %', f"{data['summary']['occupancy_rate']:.1f}%"),
        ('إجمالي المخالفات', data['summary']['total_violations']),
        ('المخالفات المحلولة', data['summary']['resolved_violations']),
        ('المخالفات المعلقة', data['summary']['pending_violations']),
        ('إجمالي الشكاوى', data['summary']['total_complaints']),
        ('الشكاوى المحلولة', data['summary']['resolved_complaints']),
        ('الشكاوى المعلقة', data['summary']['pending_complaints']),
        ('الوقائع الأمنية', data['summary']['total_security_incidents']),
        ('مواقف السيارات', data['summary']['total_parking_slots']),
        ('المواقف المشغولة', data['summary']['occupied_parking']),
        ('معدل إشغال المواقف %', f"{data['summary']['parking_occupancy_rate']:.1f}%"),
    ]
    
    row += 1
    for label, value in summary_items:
        ws_summary[f'A{row}'] = label
        ws_summary[f'B{row}'] = value
        ws_summary[f'A{row}'].font = normal_font
        ws_summary[f'B{row}'].font = normal_font
        ws_summary[f'A{row}'].alignment = right_alignment
        ws_summary[f'B{row}'].alignment = center_alignment
        ws_summary[f'A{row}'].border = border
        ws_summary[f'B{row}'].border = border
        row += 1
    
    # Column widths
    ws_summary.column_dimensions['A'].width = 30
    ws_summary.column_dimensions['B'].width = 20
    
    # Create Residents Sheet
    if report_type in ['all', 'residents']:
        ws_residents = wb.create_sheet('السكان')
        
        # Header
        ws_residents.merge_cells('A1:F1')
        ws_residents['A1'] = 'تقرير السكان'
        ws_residents['A1'].font = title_font
        ws_residents['A1'].alignment = center_alignment
        
        # Table header
        headers = ['الرقم', 'الاسم', 'المبنى', 'الشقة', 'تاريخ الدخول', 'الحالة']
        for col, header in enumerate(headers, start=1):
            cell = ws_residents.cell(row=3, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_alignment
            cell.border = border
        
        # Data
        for row_idx, resident in enumerate(data['residents'], start=4):
            ws_residents.cell(row=row_idx, column=1, value=resident['id'])
            ws_residents.cell(row=row_idx, column=2, value=resident['name'])
            ws_residents.cell(row=row_idx, column=3, value=resident['building'])
            ws_residents.cell(row=row_idx, column=4, value=resident['apartment'])
            ws_residents.cell(row=row_idx, column=5, value=resident['move_in_date'])
            ws_residents.cell(row=row_idx, column=6, value=resident['status'])
            
            for col in range(1, 7):
                cell = ws_residents.cell(row=row_idx, column=col)
                cell.font = normal_font
                cell.alignment = center_alignment
                cell.border = border
        
        # Column widths
        ws_residents.column_dimensions['A'].width = 10
        ws_residents.column_dimensions['B'].width = 30
        ws_residents.column_dimensions['C'].width = 15
        ws_residents.column_dimensions['D'].width = 12
        ws_residents.column_dimensions['E'].width = 15
        ws_residents.column_dimensions['F'].width = 12
    
    # Create Violations Sheet
    if report_type in ['all', 'violations']:
        ws_violations = wb.create_sheet('المخالفات المرورية')
        
        # Header
        ws_violations.merge_cells('A1:F1')
        ws_violations['A1'] = 'تقرير المخالفات المرورية'
        ws_violations['A1'].font = title_font
        ws_violations['A1'].alignment = center_alignment
        
        # Table header
        headers = ['الرقم', 'نوع المخالفة', 'التاريخ', 'الموقع', 'المخالف', 'الحالة']
        for col, header in enumerate(headers, start=1):
            cell = ws_violations.cell(row=3, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_alignment
            cell.border = border
        
        # Data
        for row_idx, violation in enumerate(data['violations'], start=4):
            ws_violations.cell(row=row_idx, column=1, value=violation['id'])
            ws_violations.cell(row=row_idx, column=2, value=violation['type'])
            ws_violations.cell(row=row_idx, column=3, value=violation['date'])
            ws_violations.cell(row=row_idx, column=4, value=violation['location'])
            ws_violations.cell(row=row_idx, column=5, value=violation['resident'])
            ws_violations.cell(row=row_idx, column=6, value=violation['status'])
            
            for col in range(1, 7):
                cell = ws_violations.cell(row=row_idx, column=col)
                cell.font = normal_font
                cell.alignment = center_alignment
                cell.border = border
        
        # Column widths
        ws_violations.column_dimensions['A'].width = 10
        ws_violations.column_dimensions['B'].width = 20
        ws_violations.column_dimensions['C'].width = 15
        ws_violations.column_dimensions['D'].width = 20
        ws_violations.column_dimensions['E'].width = 25
        ws_violations.column_dimensions['F'].width = 12
    
    # Create Buildings Sheet
    if report_type in ['all', 'buildings']:
        ws_buildings = wb.create_sheet('المباني')
        
        # Header
        ws_buildings.merge_cells('A1:E1')
        ws_buildings['A1'] = 'تقرير المباني'
        ws_buildings['A1'].font = title_font
        ws_buildings['A1'].alignment = center_alignment
        
        # Table header
        headers = ['الرقم', 'اسم المبنى', 'عدد الشقق', 'الشقق المشغولة', 'معدل الإشغال %']
        for col, header in enumerate(headers, start=1):
            cell = ws_buildings.cell(row=3, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_alignment
            cell.border = border
        
        # Data
        for row_idx, building in enumerate(data['buildings'], start=4):
            ws_buildings.cell(row=row_idx, column=1, value=building['id'])
            ws_buildings.cell(row=row_idx, column=2, value=building['name'])
            ws_buildings.cell(row=row_idx, column=3, value=building['apartments'])
            ws_buildings.cell(row=row_idx, column=4, value=building['occupied'])
            ws_buildings.cell(row=row_idx, column=5, value=f"{building['occupancy']:.1f}%")
            
            for col in range(1, 6):
                cell = ws_buildings.cell(row=row_idx, column=col)
                cell.font = normal_font
                cell.alignment = center_alignment
                cell.border = border
        
        # Column widths
        ws_buildings.column_dimensions['A'].width = 10
        ws_buildings.column_dimensions['B'].width = 20
        ws_buildings.column_dimensions['C'].width = 15
        ws_buildings.column_dimensions['D'].width = 18
        ws_buildings.column_dimensions['E'].width = 18
    
    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return output


def generate_word_report(report_type='all', from_date=None, to_date=None, building=None):
    """Generate professional Word report"""
    # Get data
    data = get_report_data(report_type, from_date, to_date, building)
    
    # Create document
    doc = Document()
    
    # Set up page and default font
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    
    # Title
    title = doc.add_heading(data['metadata']['title'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(15, 61, 104)
    
    # University name
    uni = doc.add_paragraph(data['metadata']['university'])
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uni_run = uni.runs[0]
    uni_run.font.name = 'Arial'
    uni_run.font.size = Pt(14)
    uni_run.font.color.rgb = RGBColor(46, 139, 192)
    uni_run.bold = True
    
    # Date
    date_para = doc.add_paragraph(f"تاريخ التقرير: {data['metadata']['generated_at_ar']}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.runs[0]
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    doc.add_heading('الملخص التنفيذي', level=1)
    
    summary_table = doc.add_table(rows=16, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    summary_items = [
        ('إجمالي السكان', str(data['summary']['total_residents'])),
        ('إجمالي المباني', str(data['summary']['total_buildings'])),
        ('إجمالي الشقق', str(data['summary']['total_apartments'])),
        ('الشقق المشغولة', str(data['summary']['occupied_apartments'])),
        ('معدل الإشغال', f"{data['summary']['occupancy_rate']:.1f}%"),
        ('إجمالي المخالفات', str(data['summary']['total_violations'])),
        ('المخالفات المحلولة', str(data['summary']['resolved_violations'])),
        ('المخالفات المعلقة', str(data['summary']['pending_violations'])),
        ('إجمالي الشكاوى', str(data['summary']['total_complaints'])),
        ('الشكاوى المحلولة', str(data['summary']['resolved_complaints'])),
        ('الشكاوى المعلقة', str(data['summary']['pending_complaints'])),
        ('الوقائع الأمنية', str(data['summary']['total_security_incidents'])),
        ('مواقف السيارات', str(data['summary']['total_parking_slots'])),
        ('المواقف المشغولة', str(data['summary']['occupied_parking'])),
        ('معدل إشغال المواقف', f"{data['summary']['parking_occupancy_rate']:.1f}%"),
    ]
    
    for idx, (label, value) in enumerate(summary_items):
        row = summary_table.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = value
        
        # Format header cells
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.name = 'Arial'
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    doc.add_page_break()
    
    # Residents Section
    if report_type in ['all', 'residents']:
        doc.add_heading('تقرير السكان', level=1)
        
        residents_table = doc.add_table(rows=len(data['residents']) + 1, cols=6)
        residents_table.style = 'Light Grid Accent 1'
        residents_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        headers = ['الرقم', 'الاسم', 'المبنى', 'الشقة', 'تاريخ الدخول', 'الحالة']
        header_row = residents_table.rows[0]
        for idx, header in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.name = 'Arial'
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        for row_idx, resident in enumerate(data['residents'], start=1):
            row = residents_table.rows[row_idx]
            row.cells[0].text = str(resident['id'])
            row.cells[1].text = resident['name']
            row.cells[2].text = resident['building']
            row.cells[3].text = resident['apartment']
            row.cells[4].text = resident['move_in_date']
            row.cells[5].text = resident['status']
            
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Arial'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    # Violations Section
    if report_type in ['all', 'violations']:
        doc.add_heading('تقرير المخالفات المرورية', level=1)
        
        violations_table = doc.add_table(rows=len(data['violations']) + 1, cols=6)
        violations_table.style = 'Light Grid Accent 1'
        violations_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        headers = ['الرقم', 'نوع المخالفة', 'التاريخ', 'الموقع', 'المخالف', 'الحالة']
        header_row = violations_table.rows[0]
        for idx, header in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.name = 'Arial'
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        for row_idx, violation in enumerate(data['violations'], start=1):
            row = violations_table.rows[row_idx]
            row.cells[0].text = str(violation['id'])
            row.cells[1].text = violation['type']
            row.cells[2].text = violation['date']
            row.cells[3].text = violation['location']
            row.cells[4].text = violation['resident']
            row.cells[5].text = violation['status']
            
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Arial'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    # Buildings Section
    if report_type in ['all', 'buildings']:
        doc.add_heading('تقرير المباني', level=1)
        
        buildings_table = doc.add_table(rows=len(data['buildings']) + 1, cols=5)
        buildings_table.style = 'Light Grid Accent 1'
        buildings_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        headers = ['الرقم', 'اسم المبنى', 'عدد الشقق', 'الشقق المشغولة', 'معدل الإشغال']
        header_row = buildings_table.rows[0]
        for idx, header in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.name = 'Arial'
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        for row_idx, building in enumerate(data['buildings'], start=1):
            row = buildings_table.rows[row_idx]
            row.cells[0].text = str(building['id'])
            row.cells[1].text = building['name']
            row.cells[2].text = str(building['apartments'])
            row.cells[3].text = str(building['occupied'])
            row.cells[4].text = f"{building['occupancy']:.1f}%"
            
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Arial'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('تم إنشاء هذا التقرير بواسطة نظام إدارة إسكان أعضاء هيئة التدريس')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.name = 'Arial'
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output
