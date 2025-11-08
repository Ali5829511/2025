"""
Housing Report Generator with PDF and Word Export
مولد تقارير الإسكان مع تصدير PDF و Word
"""

from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import database


def get_report_data():
    """Get comprehensive report data from database"""
    try:
        conn = database.get_db_connection()
        cursor = conn.cursor()
        
        report_data = {
            'metadata': {
                'report_date': datetime.now().strftime('%Y-%m-%d'),
                'report_time': datetime.now().strftime('%H:%M:%S'),
                'report_title': 'تقرير تفصيلي شامل - وحدة إسكان هيئة التدريس',
                'version': '2.0',
                'status': 'معتمد'
            },
            'overview': {},
            'buildings': {},
            'parking': {},
            'violations': {},
            'visitors': {},
            'complaints': {}
        }
        
        # Overview Statistics
        cursor.execute('SELECT COUNT(*) FROM buildings')
        report_data['overview']['total_buildings'] = cursor.fetchone()[0] or 0
        
        cursor.execute('SELECT COUNT(*) FROM apartments')
        report_data['overview']['total_units'] = cursor.fetchone()[0] or 0
        
        cursor.execute('SELECT COUNT(*) FROM residents WHERE is_active = 1')
        report_data['overview']['total_residents'] = cursor.fetchone()[0] or 0
        
        # Calculate occupancy rate
        total_units = report_data['overview']['total_units']
        occupied_units = report_data['overview']['total_residents']
        if total_units > 0:
            report_data['overview']['occupancy_rate'] = round((occupied_units / total_units) * 100, 1)
        else:
            report_data['overview']['occupancy_rate'] = 0
        
        # Buildings Statistics
        cursor.execute("SELECT COUNT(*) FROM buildings WHERE building_type = 'عمارة'")
        report_data['buildings']['apartment_count'] = cursor.fetchone()[0] or 0
        
        cursor.execute("SELECT COUNT(*) FROM buildings WHERE building_type = 'فيلا'")
        report_data['buildings']['villa_count'] = cursor.fetchone()[0] or 0
        
        cursor.execute("""
            SELECT COUNT(*) FROM buildings 
            WHERE CAST(building_number AS INTEGER) BETWEEN 1 AND 30
        """)
        report_data['buildings']['old_building_count'] = cursor.fetchone()[0] or 0
        
        cursor.execute("""
            SELECT COUNT(*) FROM buildings 
            WHERE CAST(building_number AS INTEGER) BETWEEN 53 AND 79
        """)
        report_data['buildings']['new_building_count'] = cursor.fetchone()[0] or 0
        
        # Parking Statistics
        cursor.execute('SELECT COUNT(*) FROM apartment_parking')
        report_data['parking']['total_parking'] = cursor.fetchone()[0] or 0
        
        cursor.execute("SELECT COUNT(*) FROM apartment_parking WHERE status = 'مشغول'")
        report_data['parking']['occupied_parking'] = cursor.fetchone()[0] or 0
        
        report_data['parking']['vacant_parking'] = (
            report_data['parking']['total_parking'] - 
            report_data['parking']['occupied_parking']
        )
        
        if report_data['parking']['total_parking'] > 0:
            report_data['parking']['utilization'] = round(
                (report_data['parking']['occupied_parking'] / 
                 report_data['parking']['total_parking']) * 100, 1
            )
        else:
            report_data['parking']['utilization'] = 0
        
        # Violations Statistics
        cursor.execute('SELECT COUNT(*) FROM traffic_violations')
        report_data['violations']['total_violations'] = cursor.fetchone()[0] or 0
        
        cursor.execute("""
            SELECT COUNT(*) FROM traffic_violations 
            WHERE status IN ('pending', 'open', 'مفتوحة', 'معلقة')
        """)
        report_data['violations']['open_violations'] = cursor.fetchone()[0] or 0
        
        cursor.execute("""
            SELECT COUNT(*) FROM traffic_violations 
            WHERE status IN ('resolved', 'closed', 'محلولة', 'مغلقة')
        """)
        report_data['violations']['closed_violations'] = cursor.fetchone()[0] or 0
        
        if report_data['violations']['total_violations'] > 0:
            report_data['violations']['closure_rate'] = round(
                (report_data['violations']['closed_violations'] / 
                 report_data['violations']['total_violations']) * 100, 1
            )
        else:
            report_data['violations']['closure_rate'] = 0
        
        # Visitors Statistics (placeholder - assuming table exists)
        try:
            cursor.execute("""
                SELECT COUNT(*) FROM visitors 
                WHERE DATE(visit_date) = DATE('now')
            """)
            report_data['visitors']['today_visitors'] = cursor.fetchone()[0] or 0
        except:
            report_data['visitors']['today_visitors'] = 0
        
        try:
            cursor.execute("""
                SELECT COUNT(*) FROM visitors 
                WHERE strftime('%Y-%m', visit_date) = strftime('%Y-%m', 'now')
            """)
            report_data['visitors']['month_visitors'] = cursor.fetchone()[0] or 0
        except:
            report_data['visitors']['month_visitors'] = 0
        
        # Complaints Statistics
        cursor.execute('SELECT COUNT(*) FROM complaints')
        report_data['complaints']['total_complaints'] = cursor.fetchone()[0] or 0
        
        cursor.execute("""
            SELECT COUNT(*) FROM complaints 
            WHERE status IN ('open', 'مفتوحة')
        """)
        report_data['complaints']['open_complaints'] = cursor.fetchone()[0] or 0
        
        cursor.execute("""
            SELECT COUNT(*) FROM complaints 
            WHERE status IN ('resolved', 'محلولة', 'closed', 'مغلقة')
        """)
        report_data['complaints']['closed_complaints'] = cursor.fetchone()[0] or 0
        
        if report_data['complaints']['total_complaints'] > 0:
            report_data['complaints']['resolution_rate'] = round(
                (report_data['complaints']['closed_complaints'] / 
                 report_data['complaints']['total_complaints']) * 100, 1
            )
        else:
            report_data['complaints']['resolution_rate'] = 0
        
        # Get building details
        cursor.execute("""
            SELECT 
                b.building_number,
                b.building_type,
                COUNT(DISTINCT a.id) as unit_count,
                COUNT(DISTINCT r.id) as resident_count
            FROM buildings b
            LEFT JOIN apartments a ON b.id = a.building_id
            LEFT JOIN residents r ON a.id = r.apartment_id AND r.is_active = 1
            GROUP BY b.id
            ORDER BY CAST(b.building_number AS INTEGER)
            LIMIT 10
        """)
        
        buildings_detail = []
        for row in cursor.fetchall():
            building_no, building_type, units, residents = row
            occupancy = round((residents / units * 100), 1) if units > 0 else 0
            buildings_detail.append({
                'number': building_no,
                'type': building_type,
                'units': units,
                'residents': residents,
                'occupancy': occupancy,
                'status': 'نشط'
            })
        
        report_data['buildings_detail'] = buildings_detail
        
        conn.close()
        return report_data
        
    except Exception as e:
        print(f"Error getting report data: {str(e)}")
        return None


def generate_pdf_report(report_data):
    """Generate PDF report"""
    try:
        buffer = BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=1*cm,
            leftMargin=1*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Container for PDF elements
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Create custom styles for Arabic text
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            alignment=TA_CENTER,
            fontSize=18,
            textColor=colors.HexColor('#7c3aed'),
            spaceAfter=20
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            alignment=TA_RIGHT,
            fontSize=14,
            textColor=colors.HexColor('#7c3aed'),
            spaceAfter=10
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            alignment=TA_RIGHT,
            fontSize=10
        )
        
        # Add title
        title = Paragraph(report_data['metadata']['report_title'], title_style)
        elements.append(title)
        elements.append(Spacer(1, 0.3*inch))
        
        # Add metadata table
        metadata_data = [
            ['التاريخ:', report_data['metadata']['report_date']],
            ['الوقت:', report_data['metadata']['report_time']],
            ['النسخة:', report_data['metadata']['version']],
            ['الحالة:', report_data['metadata']['status']]
        ]
        
        metadata_table = Table(metadata_data, colWidths=[2*inch, 3*inch])
        metadata_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#6b7280')),
            ('TEXTCOLOR', (1, 0), (1, -1), colors.HexColor('#1f2937')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(metadata_table)
        elements.append(Spacer(1, 0.5*inch))
        
        # Overview Section
        elements.append(Paragraph('نظرة عامة على النظام', heading_style))
        
        overview_data = [
            ['المؤشر', 'القيمة'],
            ['إجمالي المباني', str(report_data['overview']['total_buildings'])],
            ['إجمالي الوحدات', str(report_data['overview']['total_units'])],
            ['إجمالي السكان', str(report_data['overview']['total_residents'])],
            ['نسبة الإشغال', f"{report_data['overview']['occupancy_rate']}%"],
        ]
        
        overview_table = Table(overview_data, colWidths=[3*inch, 2*inch])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        elements.append(overview_table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Buildings Section
        elements.append(Paragraph('إحصائيات المباني', heading_style))
        
        buildings_data = [
            ['النوع', 'العدد'],
            ['العمارات', str(report_data['buildings']['apartment_count'])],
            ['الفلل', str(report_data['buildings']['villa_count'])],
            ['المباني القديمة', str(report_data['buildings']['old_building_count'])],
            ['المباني الجديدة', str(report_data['buildings']['new_building_count'])],
        ]
        
        buildings_table = Table(buildings_data, colWidths=[3*inch, 2*inch])
        buildings_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        elements.append(buildings_table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Parking Section
        elements.append(Paragraph('إحصائيات المواقف', heading_style))
        
        parking_data = [
            ['المؤشر', 'القيمة'],
            ['إجمالي المواقف', str(report_data['parking']['total_parking'])],
            ['المواقف المشغولة', str(report_data['parking']['occupied_parking'])],
            ['المواقف الشاغرة', str(report_data['parking']['vacant_parking'])],
            ['نسبة الاستخدام', f"{report_data['parking']['utilization']}%"],
        ]
        
        parking_table = Table(parking_data, colWidths=[3*inch, 2*inch])
        parking_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        elements.append(parking_table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Violations Section
        elements.append(Paragraph('المخالفات المرورية', heading_style))
        
        violations_data = [
            ['المؤشر', 'القيمة'],
            ['إجمالي المخالفات', str(report_data['violations']['total_violations'])],
            ['المخالفات المفتوحة', str(report_data['violations']['open_violations'])],
            ['المخالفات المغلقة', str(report_data['violations']['closed_violations'])],
            ['نسبة الإغلاق', f"{report_data['violations']['closure_rate']}%"],
        ]
        
        violations_table = Table(violations_data, colWidths=[3*inch, 2*inch])
        violations_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        elements.append(violations_table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Complaints Section
        elements.append(Paragraph('الشكاوى', heading_style))
        
        complaints_data = [
            ['المؤشر', 'القيمة'],
            ['إجمالي الشكاوى', str(report_data['complaints']['total_complaints'])],
            ['الشكاوى المفتوحة', str(report_data['complaints']['open_complaints'])],
            ['الشكاوى المغلقة', str(report_data['complaints']['closed_complaints'])],
            ['نسبة الحل', f"{report_data['complaints']['resolution_rate']}%"],
        ]
        
        complaints_table = Table(complaints_data, colWidths=[3*inch, 2*inch])
        complaints_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        elements.append(complaints_table)
        
        # Build PDF
        doc.build(elements)
        
        # Get PDF data
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return None


def generate_word_report(report_data):
    """Generate Word (DOCX) report"""
    try:
        doc = Document()
        
        # Set document direction to RTL
        section = doc.sections[0]
        
        # Add title
        title = doc.add_heading(report_data['metadata']['report_title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(124, 58, 237)
        
        # Add metadata
        metadata = doc.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        metadata.add_run(f"تاريخ التقرير: {report_data['metadata']['report_date']}\n").bold = True
        metadata.add_run(f"وقت الإنشاء: {report_data['metadata']['report_time']}\n").bold = True
        metadata.add_run(f"النسخة: {report_data['metadata']['version']}\n").bold = True
        metadata.add_run(f"الحالة: {report_data['metadata']['status']}\n").bold = True
        
        doc.add_paragraph()
        
        # Overview Section
        doc.add_heading('نظرة عامة على النظام', 1)
        overview_table = doc.add_table(rows=5, cols=2)
        overview_table.style = 'Light Grid Accent 1'
        
        overview_data = [
            ('إجمالي المباني', str(report_data['overview']['total_buildings'])),
            ('إجمالي الوحدات', str(report_data['overview']['total_units'])),
            ('إجمالي السكان', str(report_data['overview']['total_residents'])),
            ('نسبة الإشغال', f"{report_data['overview']['occupancy_rate']}%"),
        ]
        
        header_cells = overview_table.rows[0].cells
        header_cells[0].text = 'المؤشر'
        header_cells[1].text = 'القيمة'
        
        for i, (label, value) in enumerate(overview_data, start=1):
            row_cells = overview_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
        
        doc.add_paragraph()
        
        # Buildings Section
        doc.add_heading('إحصائيات المباني', 1)
        buildings_table = doc.add_table(rows=5, cols=2)
        buildings_table.style = 'Light Grid Accent 1'
        
        buildings_data = [
            ('العمارات', str(report_data['buildings']['apartment_count'])),
            ('الفلل', str(report_data['buildings']['villa_count'])),
            ('المباني القديمة', str(report_data['buildings']['old_building_count'])),
            ('المباني الجديدة', str(report_data['buildings']['new_building_count'])),
        ]
        
        header_cells = buildings_table.rows[0].cells
        header_cells[0].text = 'النوع'
        header_cells[1].text = 'العدد'
        
        for i, (label, value) in enumerate(buildings_data, start=1):
            row_cells = buildings_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
        
        doc.add_paragraph()
        
        # Parking Section
        doc.add_heading('إحصائيات المواقف', 1)
        parking_table = doc.add_table(rows=5, cols=2)
        parking_table.style = 'Light Grid Accent 1'
        
        parking_data = [
            ('إجمالي المواقف', str(report_data['parking']['total_parking'])),
            ('المواقف المشغولة', str(report_data['parking']['occupied_parking'])),
            ('المواقف الشاغرة', str(report_data['parking']['vacant_parking'])),
            ('نسبة الاستخدام', f"{report_data['parking']['utilization']}%"),
        ]
        
        header_cells = parking_table.rows[0].cells
        header_cells[0].text = 'المؤشر'
        header_cells[1].text = 'القيمة'
        
        for i, (label, value) in enumerate(parking_data, start=1):
            row_cells = parking_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
        
        doc.add_paragraph()
        
        # Violations Section
        doc.add_heading('المخالفات المرورية', 1)
        violations_table = doc.add_table(rows=5, cols=2)
        violations_table.style = 'Light Grid Accent 1'
        
        violations_data = [
            ('إجمالي المخالفات', str(report_data['violations']['total_violations'])),
            ('المخالفات المفتوحة', str(report_data['violations']['open_violations'])),
            ('المخالفات المغلقة', str(report_data['violations']['closed_violations'])),
            ('نسبة الإغلاق', f"{report_data['violations']['closure_rate']}%"),
        ]
        
        header_cells = violations_table.rows[0].cells
        header_cells[0].text = 'المؤشر'
        header_cells[1].text = 'القيمة'
        
        for i, (label, value) in enumerate(violations_data, start=1):
            row_cells = violations_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
        
        doc.add_paragraph()
        
        # Complaints Section
        doc.add_heading('الشكاوى', 1)
        complaints_table = doc.add_table(rows=5, cols=2)
        complaints_table.style = 'Light Grid Accent 1'
        
        complaints_data = [
            ('إجمالي الشكاوى', str(report_data['complaints']['total_complaints'])),
            ('الشكاوى المفتوحة', str(report_data['complaints']['open_complaints'])),
            ('الشكاوى المغلقة', str(report_data['complaints']['closed_complaints'])),
            ('نسبة الحل', f"{report_data['complaints']['resolution_rate']}%"),
        ]
        
        header_cells = complaints_table.rows[0].cells
        header_cells[0].text = 'المؤشر'
        header_cells[1].text = 'القيمة'
        
        for i, (label, value) in enumerate(complaints_data, start=1):
            row_cells = complaints_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
        
        doc.add_paragraph()
        
        # Summary Section
        doc.add_heading('الخلاصة والتوصيات', 1)
        summary_para = doc.add_paragraph()
        summary_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        summary_para.add_run(f"• نسبة الإشغال الإجمالية: {report_data['overview']['occupancy_rate']}%\n")
        vacant_units = report_data['overview']['total_units'] - report_data['overview']['total_residents']
        summary_para.add_run(f"• عدد الوحدات الشاغرة: {vacant_units} وحدة\n")
        summary_para.add_run(f"• المواقف المتاحة: {report_data['parking']['vacant_parking']} موقف\n")
        summary_para.add_run(f"• المخالفات المفتوحة: {report_data['violations']['open_violations']} مخالفة\n")
        summary_para.add_run(f"• الشكاوى المفتوحة: {report_data['complaints']['open_complaints']} شكوى\n")
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        print(f"Error generating Word report: {str(e)}")
        return None
